from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
from io import BytesIO
import re
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import docx


def export_events_to_words(events):
    doc: Document = Document()
    for event in events:
        heading = doc.add_heading(event["event_name"], level=1)
        heading_run = heading.runs[0]
        heading_run.bold = True
        heading_run.italic = True
        heading_run.font_size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        date_str = event["date_created"].strftime("%d.%m.%Y")
        content = "Ngày " + date_str + "\n" + event["event_content"]
        doc_content = doc.add_paragraph(content)
        doc_content.add_run("\nNguồn tin \n").italic = True
        for source in event["source_list"]:
            doc_content.add_run(f"+ {source}\n").italic = True
    buff = BytesIO()
    doc.save(buff)
    buff.seek(0)
    return buff


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def insert_html_to_docx(doc, html):
    html = html.split("<")
    html = [html[0]] + ["<" + l for l in html[1:]]
    tags = []
    # doc = Document()
    p = doc.add_paragraph()
    for run in html:
        tag_change = re.match("(?:<)(.*?)(?:>)", run)
        if tag_change != None:
            tag_strip = tag_change.group(0)
            tag_change = tag_change.group(1)
            if tag_change.startswith("/"):
                if tag_change.startswith("/a"):
                    tag_change = next(tag for tag in tags if tag.startswith("a "))
                tag_change = tag_change.strip("/")
                tags.remove(tag_change)
            else:
                tags.append(tag_change)
        else:
            tag_strip = ""
        hyperlink = [tag for tag in tags if tag.startswith("a ")]
        if run.startswith("<"):
            run = run.replace(tag_strip, "")
            if hyperlink:
                hyperlink = hyperlink[0]
                hyperlink = re.match('.*?(?:href=")(.*?)(?:").*?', hyperlink).group(1)
                add_hyperlink(p, run, hyperlink)
            else:
                runner = p.add_run(run)
                if "b" in tags:
                    runner.bold = True
                if "u" in tags:
                    runner.underline = True
                if "i" in tags:
                    runner.italic = True
                if "H1" in tags:
                    runner.font.size = Pt(24)
        else:
            p.add_run(run)


def add_hyperlink_into_run(paragraph, run, url):
    # runs = paragraph.runs
    # for i in range(len(runs)):
    #     if runs[i].text == run.text:
    #         break
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )
    hyperlink.append(run._r)
    paragraph._p.insert(2, hyperlink)
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    run.italic = True


def export_news_to_words(news):
    doc: Document = Document()
    for event in news:
        heading = doc.add_heading("-" + event["data:title"], level=1)
        heading_run = heading.runs[0]
        heading_run.bold = True
        heading_run.italic = True
        heading_run.font_size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        date_str = event["pub_date"].strftime("%d.%m.%Y")
        doc.add_paragraph(f"Ngày {date_str}")
        # content = "Ngày " + date_str + "\n" + event["event_content"]
        # doc_content = doc.add_paragraph(content)
        # insert_html_to_docx(doc, event["data:content"])
        doc_content = doc.add_paragraph(event["data:content"])
        doc_content = doc.add_paragraph("\n")
        doc_content.add_run("Nguồn tin \n").italic = True
        doc_run = doc_content.add_run(
            f"+ ({event['source_host_name']}) {event['data:title']}\n"
        )
        add_hyperlink_into_run(doc_content, doc_run, event["data:url"])
    buff = BytesIO()
    doc.save(buff)
    buff.seek(0)
    return buff


def export_social_word(news_list, platform: str = ""):
    doc: Document = Document()
    for news in news_list:
        header = news.get("header")
        content = news.get("content")
        create_at = datetime.strptime(news.get("created_at"), "%Y/%m/%d %H:%M:%S")
        heading = doc.add_heading(header, level=1)
        heading_run = heading.runs[0]
        heading_run.bold = True
        heading_run.italic = True
        heading_run.font_size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        date_str = create_at.strftime("%d.%m.%Y")
        content = "Ngày " + date_str + "\n\n" + content.replace("\nSee translation", "")
        doc_content = doc.add_paragraph(content)
        doc_source = doc.add_paragraph()
        # doc_source.add_run("Nguồn tin \n").italic = True
        # doc_source.add_run(f"+ {news.get('id')}\n").italic = True
    buff = BytesIO()
    doc.save(buff)
    buff.seek(0)
    return buff
